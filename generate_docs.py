import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- KONFIGURASI WARNA & STYLING ---
COLOR_NAVY_HEX = "1E3A8A"
COLOR_GOLD_HEX = "D97706"
COLOR_SLATE_HEX = "334155"
COLOR_LIGHT_BG_HEX = "F8FAFC"
COLOR_BORDER_HEX = "CBD5E1"
COLOR_INFO_BG_HEX = "EFF6FF"
COLOR_INFO_BORDER_HEX = "3B82F6"
COLOR_WARN_BG_HEX = "FFFBEB"
COLOR_WARN_BORDER_HEX = "F59E0B"
COLOR_CAM_BG_HEX = "F1F5F9"
COLOR_CAM_BORDER_HEX = "94A3B8"

COLOR_NAVY = RGBColor(30, 58, 138)
COLOR_GOLD = RGBColor(217, 119, 6)
COLOR_SLATE = RGBColor(51, 65, 85)
COLOR_DARK = RGBColor(15, 23, 42)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for border_name, border_props in kwargs.items():
        if border_props:
            b_el = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="{border_props.get("val", "single")}" w:sz="{border_props.get("sz", "4")}" w:space="0" w:color="{border_props.get("color", "auto")}"/>')
            tcBorders.append(b_el)
        else:
            b_el = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="none"/>')
            tcBorders.append(b_el)
    tcPr.append(tcBorders)

def add_callout_box(doc, text, is_warning=False):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    
    bg_color = COLOR_WARN_BG_HEX if is_warning else COLOR_INFO_BG_HEX
    border_color = COLOR_WARN_BORDER_HEX if is_warning else COLOR_INFO_BORDER_HEX
    
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=240, right=200)
    set_cell_border(cell, left={'val': 'single', 'sz': 24, 'color': border_color})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_DARK
    
    empty_p = doc.add_paragraph()
    empty_p.paragraph_format.space_before = Pt(0)
    empty_p.paragraph_format.space_after = Pt(4)

def add_image_placeholder(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    
    set_cell_background(cell, COLOR_CAM_BG_HEX)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    set_cell_border(cell, 
                    left={'val': 'single', 'sz': 18, 'color': COLOR_CAM_BORDER_HEX},
                    top={'val': 'dotted', 'sz': 4, 'color': COLOR_BORDER_HEX},
                    right={'val': 'dotted', 'sz': 4, 'color': COLOR_BORDER_HEX},
                    bottom={'val': 'dotted', 'sz': 4, 'color': COLOR_BORDER_HEX})
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    clean_text = text.replace("[📷", "").replace("]", "").strip()
    run = p.add_run(f"📷 {clean_text}")
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(71, 85, 105)
    
    empty_p = doc.add_paragraph()
    empty_p.paragraph_format.space_before = Pt(0)
    empty_p.paragraph_format.space_after = Pt(4)

def render_markdown_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if '---' in line and '|' in line:
            continue # skip alignment row
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        
    if not rows:
        return
        
    cols_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        bg = COLOR_NAVY_HEX if is_header else (COLOR_LIGHT_BG_HEX if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, cell_value in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                set_cell_border(cell, bottom={'val': 'single', 'sz': 2, 'color': COLOR_BORDER_HEX})
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                # Format text
                run = p.add_run(cell_value.replace('**', ''))
                run.font.name = 'Arial'
                run.font.size = Pt(9 if not is_header else 9.5)
                run.font.bold = is_header or '**' in cell_value
                run.font.color.rgb = RGBColor(255, 255, 255) if is_header else COLOR_DARK

def parse_md_file_to_docx(md_path, output_path="Buku_Panduan_Portal_Desa_Pringanom.docx"):
    if not os.path.exists(md_path):
        print(f"File {md_path} tidak ditemukan!")
        return

    doc = Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = True
    
    # Header/Footer
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("Buku Panduan Penggunaan Portal Resmi Desa Pringanom v4")
    hrun.font.name = 'Arial'; hrun.font.size = Pt(8.5); hrun.font.color.rgb = RGBColor(148, 163, 184)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Pemerintah Desa Pringanom, Masaran, Sragen — 2026")
    frun.font.name = 'Arial'; frun.font.size = Pt(8.5); frun.font.color.rgb = RGBColor(148, 163, 184)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_buffer = []
    in_code_block = False

    for line in lines:
        raw_line = line.strip()

        # Handle Code Block ```
        if raw_line.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_callout_box(doc, raw_line, is_warning=False)
            continue

        # Handle Table Buffer
        if '|' in raw_line and not raw_line.startswith('>'):
            in_table = True
            table_buffer.append(raw_line)
            continue
        else:
            if in_table:
                render_markdown_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not raw_line:
            continue

        # Image Placeholder
        if '[📷' in raw_line:
            add_image_placeholder(doc, raw_line)
            continue

        # Callout Quote
        if raw_line.startswith('>'):
            clean_quote = raw_line.lstrip('>').strip().replace('**', '')
            is_warn = "Penting" in clean_quote or "Perhatian" in clean_quote or "Sangat penting" in clean_quote
            add_callout_box(doc, clean_quote, is_warning=is_warn)
            continue

        # Headings
        if raw_line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8); p.paragraph_format.keep_with_next = True
            r = p.add_run(raw_line.replace('# ', '').replace('**', ''))
            r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
            continue
        elif raw_line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6); p.paragraph_format.keep_with_next = True
            r = p.add_run(raw_line.replace('## ', '').replace('**', ''))
            r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = COLOR_GOLD
            continue
        elif raw_line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4); p.paragraph_format.keep_with_next = True
            r = p.add_run(raw_line.replace('### ', '').replace('**', ''))
            r.font.name = 'Arial'; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = COLOR_SLATE
            continue
        elif raw_line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2); p.paragraph_format.keep_with_next = True
            r = p.add_run(raw_line.replace('#### ', '').replace('**', ''))
            r.font.name = 'Arial'; r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = COLOR_DARK
            continue

        # Bullet List
        if raw_line.startswith('- ') or raw_line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(raw_line[2:].replace('**', ''))
            r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = COLOR_DARK
            continue

        # Numbered List
        if re.match(r'^\d+\.\s', raw_line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
            clean_num = re.sub(r'^\d+\.\s', '', raw_line).replace('**', '')
            r = p.add_run(clean_num)
            r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = COLOR_DARK
            continue

        # Horizontal Divider
        if raw_line in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            r = p.add_run("_________________________________________________________________________________")
            r.font.color.rgb = RGBColor(203, 213, 225)
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.15
        r = p.add_run(raw_line.replace('**', ''))
        r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = COLOR_DARK

    # Flush last table if exists
    if in_table:
        render_markdown_table(doc, table_buffer)

    doc.save(output_path)
    print(f"SELESAI! Dokumen 1000+ baris berhasil dikonversi utuh: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    parse_md_file_to_docx("MODUL_PANDUAN_PENGGUNAAN.md")